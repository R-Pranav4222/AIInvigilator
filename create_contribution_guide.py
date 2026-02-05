from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Title
title = doc.add_heading('AIInvigilator - GitHub Contribution Guide', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('Guide for Teammates to Contribute to the Project')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(14)
subtitle_format.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()

# Step 1
doc.add_heading('Step 1: Clone the Repository (First Time Only)', 1)
doc.add_paragraph('Navigate to where you want the project and clone the repository:')
p = doc.add_paragraph('cd C:\\Users\\YourName\\Documents', style='Intense Quote')
p = doc.add_paragraph('git clone https://github.com/SarangSanthosh/AIInvigilator.git', style='Intense Quote')
p = doc.add_paragraph('cd AIInvigilator', style='Intense Quote')

# Step 2
doc.add_heading('Step 2: Set Up the Project', 1)
doc.add_paragraph('Install dependencies and set up the database:')
p = doc.add_paragraph('pip install -r requirements.txt', style='Intense Quote')
p = doc.add_paragraph('# Create .env file with database credentials', style='Intense Quote')
p = doc.add_paragraph('python manage.py migrate', style='Intense Quote')
p = doc.add_paragraph('python manage.py createsuperuser  # Optional', style='Intense Quote')

# Step 3
doc.add_heading('Step 3: Create a New Branch for Your Work', 1)
important = doc.add_paragraph()
important.add_run('Important: ').bold = True
important.add_run('Always work on a separate branch, not directly on main!')
doc.add_paragraph()
p = doc.add_paragraph('git checkout -b feature/your-feature-name', style='Intense Quote')
doc.add_paragraph('Example branch names:')
doc.add_paragraph('• git checkout -b feature/add-email-notifications', style='List Bullet')
doc.add_paragraph('• git checkout -b bugfix/fix-mobile-detection', style='List Bullet')
doc.add_paragraph('• git checkout -b feature/improve-ui', style='List Bullet')

# Step 4
doc.add_heading('Step 4: Make Changes', 1)
doc.add_paragraph('Your teammate can now:')
doc.add_paragraph('• Edit files', style='List Bullet')
doc.add_paragraph('• Add new features', style='List Bullet')
doc.add_paragraph('• Fix bugs', style='List Bullet')
doc.add_paragraph('• Test their changes locally', style='List Bullet')
doc.add_paragraph()
doc.add_paragraph('Test the application:')
p = doc.add_paragraph('python manage.py runserver', style='Intense Quote')
p = doc.add_paragraph('cd ML', style='Intense Quote')
p = doc.add_paragraph('python front.py', style='Intense Quote')

# Step 5
doc.add_heading('Step 5: Check What Changed', 1)
p = doc.add_paragraph('git status  # See which files were modified', style='Intense Quote')
p = doc.add_paragraph('git diff  # See detailed changes in files', style='Intense Quote')

# Step 6
doc.add_heading('Step 6: Stage and Commit Changes', 1)
p = doc.add_paragraph('# Stage specific files', style='Intense Quote')
p = doc.add_paragraph('git add path/to/file1.py', style='Intense Quote')
p = doc.add_paragraph('git add path/to/file2.html', style='Intense Quote')
p = doc.add_paragraph('', style='Intense Quote')
p = doc.add_paragraph('# Or stage all changes', style='Intense Quote')
p = doc.add_paragraph('git add .', style='Intense Quote')
p = doc.add_paragraph('', style='Intense Quote')
p = doc.add_paragraph('# Commit with a descriptive message', style='Intense Quote')
p = doc.add_paragraph('git commit -m "Add email notification feature for malpractice detection"', style='Intense Quote')
doc.add_paragraph()
doc.add_paragraph('Example commit messages:')
doc.add_paragraph('• "Fix turning back detection threshold"', style='List Bullet')
doc.add_paragraph('• "Improve passing paper detection algorithm"', style='List Bullet')
doc.add_paragraph('• "Update UI for malpractice log page"', style='List Bullet')

# Step 7
doc.add_heading('Step 7: Pull Latest Changes from Main', 1)
doc.add_paragraph('Before pushing, always pull the latest changes to avoid conflicts:')
p = doc.add_paragraph('git checkout main', style='Intense Quote')
p = doc.add_paragraph('git pull origin main', style='Intense Quote')
p = doc.add_paragraph('git checkout feature/your-feature-name', style='Intense Quote')
p = doc.add_paragraph('git merge main', style='Intense Quote')
doc.add_paragraph()
doc.add_paragraph('If there are merge conflicts, resolve them:')
doc.add_paragraph('1. Open conflicting files in VS Code', style='List Number')
doc.add_paragraph('2. Choose which changes to keep', style='List Number')
doc.add_paragraph('3. Save the files', style='List Number')
doc.add_paragraph('4. Stage and commit the resolved files', style='List Number')
p = doc.add_paragraph('git add .', style='Intense Quote')
p = doc.add_paragraph('git commit -m "Resolve merge conflicts"', style='Intense Quote')

# Step 8
doc.add_heading('Step 8: Push Changes to GitHub', 1)
p = doc.add_paragraph('git push origin feature/your-feature-name', style='Intense Quote')

# Step 9
doc.add_heading('Step 9: Create a Pull Request (PR)', 1)
doc.add_paragraph('1. Go to GitHub repository: https://github.com/SarangSanthosh/AIInvigilator', style='List Number')
doc.add_paragraph('2. Click the "Compare & pull request" button', style='List Number')
doc.add_paragraph('3. Fill in the PR details:', style='List Number')
doc.add_paragraph('   • Title: Brief description (e.g., "Add email notifications")', style='List Bullet')
doc.add_paragraph('   • Description: Explain what you changed and why', style='List Bullet')
doc.add_paragraph('4. Click "Create pull request"', style='List Number')

# Step 10
doc.add_heading('Step 10: Review and Merge', 1)
doc.add_paragraph('Project owner will:')
doc.add_paragraph('1. Review the pull request on GitHub', style='List Number')
doc.add_paragraph('2. Check the changes', style='List Number')
doc.add_paragraph('3. Test if needed', style='List Number')
doc.add_paragraph('4. Click "Merge pull request" if everything looks good', style='List Number')
doc.add_paragraph('5. Delete the feature branch (optional)', style='List Number')

# Step 11
doc.add_heading('Step 11: Update Your Local Repository After Merge', 1)
p = doc.add_paragraph('git checkout main', style='Intense Quote')
p = doc.add_paragraph('git pull origin main', style='Intense Quote')
p = doc.add_paragraph('git branch -d feature/your-feature-name  # Delete old branch', style='Intense Quote')

# Add page break
doc.add_page_break()

# Quick Reference
doc.add_heading('Quick Reference Cheat Sheet', 1)

doc.add_heading('First Time Setup', 2)
p = doc.add_paragraph('git clone https://github.com/SarangSanthosh/AIInvigilator.git', style='Intense Quote')
p = doc.add_paragraph('cd AIInvigilator', style='Intense Quote')
p = doc.add_paragraph('pip install -r requirements.txt', style='Intense Quote')

doc.add_heading('Start Working on New Feature', 2)
p = doc.add_paragraph('git checkout -b feature/my-feature', style='Intense Quote')
p = doc.add_paragraph('# ... make changes ...', style='Intense Quote')
p = doc.add_paragraph('git add .', style='Intense Quote')
p = doc.add_paragraph('git commit -m "Description of changes"', style='Intense Quote')

doc.add_heading('Before Pushing', 2)
p = doc.add_paragraph('git checkout main', style='Intense Quote')
p = doc.add_paragraph('git pull origin main', style='Intense Quote')
p = doc.add_paragraph('git checkout feature/my-feature', style='Intense Quote')
p = doc.add_paragraph('git merge main', style='Intense Quote')

doc.add_heading('Push Changes', 2)
p = doc.add_paragraph('git push origin feature/my-feature', style='Intense Quote')
p = doc.add_paragraph('# Then create Pull Request on GitHub', style='Intense Quote')

doc.add_heading('After PR is Merged', 2)
p = doc.add_paragraph('git checkout main', style='Intense Quote')
p = doc.add_paragraph('git pull origin main', style='Intense Quote')

# Add page break
doc.add_page_break()

# Common Git Commands
doc.add_heading('Common Git Commands', 1)

doc.add_paragraph('Check current branch:')
p = doc.add_paragraph('git branch', style='Intense Quote')

doc.add_paragraph('Switch branches:')
p = doc.add_paragraph('git checkout branch-name', style='Intense Quote')

doc.add_paragraph('See commit history:')
p = doc.add_paragraph('git log --oneline', style='Intense Quote')

doc.add_paragraph('Undo last commit (keep changes):')
p = doc.add_paragraph('git reset --soft HEAD~1', style='Intense Quote')

doc.add_paragraph('Discard all local changes:')
p = doc.add_paragraph('git reset --hard', style='Intense Quote')

doc.add_paragraph('Update from remote:')
p = doc.add_paragraph('git fetch origin', style='Intense Quote')
p = doc.add_paragraph('git pull origin main', style='Intense Quote')

doc.add_paragraph('See remote repository URL:')
p = doc.add_paragraph('git remote -v', style='Intense Quote')

# Best Practices
doc.add_heading('Best Practices for Your Team', 1)
doc.add_paragraph('1. Always work on feature branches, never directly on main', style='List Number')
doc.add_paragraph('2. Pull latest changes before starting new work', style='List Number')
doc.add_paragraph('3. Commit frequently with clear messages', style='List Number')
doc.add_paragraph('4. Test your changes before pushing', style='List Number')
doc.add_paragraph('5. Create Pull Requests for code review', style='List Number')
doc.add_paragraph('6. Write descriptive commit messages', style='List Number')
doc.add_paragraph('7. Keep PRs small and focused (one feature per PR)', style='List Number')
doc.add_paragraph('8. Communicate with team about major changes', style='List Number')

# Example Workflow
doc.add_heading('Example Workflow', 1)

doc.add_heading('Day 1: Start New Feature', 2)
p = doc.add_paragraph('git checkout main', style='Intense Quote')
p = doc.add_paragraph('git pull origin main', style='Intense Quote')
p = doc.add_paragraph('git checkout -b feature/add-sms-alerts', style='Intense Quote')
p = doc.add_paragraph('# ... edit files ...', style='Intense Quote')
p = doc.add_paragraph('git add app/views.py', style='Intense Quote')
p = doc.add_paragraph('git commit -m "Add SMS alert functionality"', style='Intense Quote')
p = doc.add_paragraph('git push origin feature/add-sms-alerts', style='Intense Quote')
p = doc.add_paragraph('# Create PR on GitHub', style='Intense Quote')

doc.add_heading('Day 2: Continue Work After Feedback', 2)
p = doc.add_paragraph('git checkout feature/add-sms-alerts', style='Intense Quote')
p = doc.add_paragraph('# ... make more changes ...', style='Intense Quote')
p = doc.add_paragraph('git add app/views.py templates/alerts.html', style='Intense Quote')
p = doc.add_paragraph('git commit -m "Update SMS alert UI based on review"', style='Intense Quote')
p = doc.add_paragraph('git push origin feature/add-sms-alerts', style='Intense Quote')
p = doc.add_paragraph('# Update existing PR automatically', style='Intense Quote')

doc.add_heading('After PR is Merged', 2)
p = doc.add_paragraph('git checkout main', style='Intense Quote')
p = doc.add_paragraph('git pull origin main', style='Intense Quote')
p = doc.add_paragraph('git branch -d feature/add-sms-alerts', style='Intense Quote')

# Footer
doc.add_paragraph()
doc.add_paragraph()
footer = doc.add_paragraph('This guide ensures your teammates can contribute safely without overwriting each other\'s work! 🚀')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_format = footer.runs[0]
footer_format.font.size = Pt(12)
footer_format.font.color.rgb = RGBColor(0, 128, 0)

# Save document
doc.save('GitHub_Contribution_Guide.docx')
print("Word document created successfully: GitHub_Contribution_Guide.docx")
