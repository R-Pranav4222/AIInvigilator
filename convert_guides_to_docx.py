"""
Convert all 4 Study Guide Markdown files to .docx format.
Handles: headings, bold/italic/code spans, code blocks, tables, bullet lists,
         horizontal rules, and LaTeX math (inline $...$ and block $$...$$).
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree
import latex2mathml.converter

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# ──────────────────────── Math (LaTeX → OMML) ────────────────────────

MML2OMML_XSL = r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL"
_xslt_transform = None

def _get_xslt():
    """Lazy-load the MML2OMML XSLT transform."""
    global _xslt_transform
    if _xslt_transform is None:
        with open(MML2OMML_XSL, 'rb') as f:
            xsl_doc = etree.parse(f)
        _xslt_transform = etree.XSLT(xsl_doc)
    return _xslt_transform


def latex_to_omml(latex_str):
    """Convert a LaTeX math string to an OMML element for python-docx.
    Returns an lxml element (<m:oMath> or <m:oMathPara>) or None on failure."""
    try:
        # Step 1: LaTeX → MathML
        mathml_str = latex2mathml.converter.convert(latex_str)

        # Step 2: Parse MathML
        mathml_tree = etree.fromstring(mathml_str.encode('utf-8'))

        # Step 3: MathML → OMML via XSLT
        xslt = _get_xslt()
        omml_tree = xslt(mathml_tree)

        # The transform returns a tree; get the root element
        omml_root = omml_tree.getroot()

        # Return the first oMath or oMathPara element
        nsmap = {'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
        for tag in ['m:oMathPara', 'm:oMath']:
            found = omml_root.findall(f'.//{tag}', nsmap)
            if found:
                return found[0]
        # Fallback: return root itself if it's an oMath element
        if omml_root.tag.endswith('}oMathPara') or omml_root.tag.endswith('}oMath'):
            return omml_root
        # If root has children, return first child
        if len(omml_root):
            return omml_root[0]
        return omml_root
    except Exception as e:
        print(f"    [WARN] Math conversion failed for: {latex_str[:60]}... → {e}")
        return None


def add_block_math(doc, latex_str):
    """Add a display-style (block) math equation as a native Word equation."""
    omml_elem = latex_to_omml(latex_str)
    if omml_elem is not None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p._p.append(omml_elem)
    else:
        # Fallback: render as styled text
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(latex_str)
        run.font.name = 'Cambria Math'
        run.font.size = Pt(11)
        run.italic = True


def add_inline_math_to_paragraph(paragraph, latex_str):
    """Insert an inline math equation into an existing paragraph."""
    omml_elem = latex_to_omml(latex_str)
    if omml_elem is not None:
        # For inline, we want oMath not oMathPara
        nsmap = {'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
        # If we got an oMathPara, extract the oMath inside it
        if omml_elem.tag.endswith('}oMathPara'):
            inner = omml_elem.findall('.//m:oMath', nsmap)
            if inner:
                omml_elem = inner[0]
        paragraph._p.append(omml_elem)
    else:
        # Fallback: render as italic Cambria Math
        run = paragraph.add_run(latex_str)
        run.font.name = 'Cambria Math'
        run.font.size = Pt(10)
        run.italic = True

# ──────────────────────── Style helpers ────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_text(paragraph, text):
    """Parse inline markdown (bold, italic, inline code, links, inline math) and add runs."""
    # Pattern order matters: bold+italic first, then bold, italic, code, links, inline math
    patterns = [
        (r'\*\*\*(.+?)\*\*\*', {'bold': True, 'italic': True}),
        (r'\*\*(.+?)\*\*', {'bold': True}),
        (r'\*(.+?)\*', {'italic': True}),
        (r'`([^`]+)`', {'code': True}),
        (r'\[([^\]]+)\]\([^\)]+\)', {'link': True}),
        (r'\$([^\$]+?)\$', {'math': True}),
    ]

    # Split text by all inline patterns (including inline math $...$)
    combined = (r'(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*[^*]+?\*|`[^`]+`'
                r'|\[[^\]]+\]\([^\)]+\)'
                r'|\$[^\$]+?\$)')
    parts = re.split(combined, text)

    for part in parts:
        if not part:
            continue

        matched = False
        for pattern, fmt in patterns:
            m = re.fullmatch(pattern, part)
            if m:
                if fmt.get('math'):
                    # Insert native Word inline equation
                    add_inline_math_to_paragraph(paragraph, m.group(1))
                else:
                    run = paragraph.add_run(m.group(1))
                    run.font.size = paragraph.style.font.size or Pt(11)
                    if fmt.get('bold'):
                        run.bold = True
                    if fmt.get('italic'):
                        run.italic = True
                    if fmt.get('code'):
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9.5)
                        run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
                    if fmt.get('link'):
                        run.font.color.rgb = RGBColor(0x06, 0x45, 0xAD)
                matched = True
                break

        if not matched:
            run = paragraph.add_run(part)


def add_code_block(doc, lines):
    """Add a shaded code block as a single paragraph with monospace font."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    code_text = '\n'.join(lines)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)

    # Light grey background on paragraph
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="F2F2F2"/>')
    pPr.append(shading)


def add_table(doc, header_row, data_rows):
    """Add a formatted table to the document."""
    if not header_row:
        return

    cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    # Header row
    for i, hdr in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(hdr.strip())
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "2E74B5")

    # Data rows
    for r_idx, row_data in enumerate(data_rows):
        for c_idx in range(cols):
            cell_text = row_data[c_idx].strip() if c_idx < len(row_data) else ''
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            add_formatted_text(p, cell_text)
            p.style = doc.styles['Normal']
            for run in p.runs:
                if run.font.size is None:
                    run.font.size = Pt(9.5)
            # Alternating row shading
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F7FB")

    doc.add_paragraph('')  # spacer


# ──────────────────────── Main converter ────────────────────────

def convert_md_to_docx(md_path, docx_path):
    """Convert a single markdown file to .docx."""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    # Set heading styles
    for level in range(1, 5):
        try:
            hs = doc.styles[f'Heading {level}']
            hs.font.name = 'Calibri'
            hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6D)
            if level == 1:
                hs.font.size = Pt(22)
            elif level == 2:
                hs.font.size = Pt(16)
            elif level == 3:
                hs.font.size = Pt(13)
            elif level == 4:
                hs.font.size = Pt(11.5)
        except KeyError:
            pass

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')

        # ─── Blank line ───
        if not line.strip():
            i += 1
            continue

        # ─── Horizontal rule ───
        if re.match(r'^-{3,}$', line.strip()) or re.match(r'^\*{3,}$', line.strip()):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Add a thin line via border
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue

        # ─── Code block ───
        if line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i].rstrip('\n'))
                i += 1
            i += 1  # skip closing ```
            add_code_block(doc, code_lines)
            continue

        # ─── Block math ($$...$$) ───
        if line.strip().startswith('$$'):
            # Could be single-line $$...$$ or multi-line
            stripped = line.strip()
            if stripped.endswith('$$') and len(stripped) > 4:
                # Single-line: $$P = ...$$
                latex = stripped[2:-2].strip()
                add_block_math(doc, latex)
                i += 1
            else:
                # Multi-line block
                math_lines = []
                if len(stripped) > 2:
                    math_lines.append(stripped[2:])
                i += 1
                while i < len(lines) and '$$' not in lines[i].strip():
                    math_lines.append(lines[i].rstrip('\n'))
                    i += 1
                if i < len(lines):
                    closing = lines[i].strip()
                    if closing.endswith('$$'):
                        rest = closing[:-2].strip()
                        if rest:
                            math_lines.append(rest)
                    i += 1  # skip closing $$
                latex = ' '.join(l.strip() for l in math_lines)
                add_block_math(doc, latex)
            continue

        # ─── Table ───
        if '|' in line and not line.strip().startswith('```'):
            table_lines = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip():
                table_lines.append(lines[i].strip())
                i += 1

            if len(table_lines) >= 2:
                def parse_row(row_str):
                    cells = row_str.split('|')
                    # Remove empty first/last from leading/trailing |
                    if cells and not cells[0].strip():
                        cells = cells[1:]
                    if cells and not cells[-1].strip():
                        cells = cells[:-1]
                    return [c.strip() for c in cells]

                header = parse_row(table_lines[0])
                # Skip separator row (row with ---)
                data_start = 1
                if data_start < len(table_lines) and re.match(r'^[\|\s\-:]+$', table_lines[data_start]):
                    data_start = 2

                data = [parse_row(r) for r in table_lines[data_start:]]
                add_table(doc, header, data)
            continue

        # ─── Headings ───
        heading_match = re.match(r'^(#{1,4})\s+(.*)', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            p = doc.add_heading('', level=level)
            add_formatted_text(p, text)
            i += 1
            continue

        # ─── Bullet / numbered list ───
        list_match = re.match(r'^(\s*)([-*+]|\d+[.\)])\s+(.*)', line)
        if list_match:
            indent_level = len(list_match.group(1)) // 2
            marker = list_match.group(2)
            text = list_match.group(3)

            if re.match(r'\d+[.\)]', marker):
                p = doc.add_paragraph('', style='List Number')
            else:
                p = doc.add_paragraph('', style='List Bullet')

            if indent_level > 0:
                p.paragraph_format.left_indent = Cm(1.27 * (indent_level + 1))

            add_formatted_text(p, text)
            i += 1
            continue

        # ─── Regular paragraph ───
        p = doc.add_paragraph('')
        add_formatted_text(p, line)
        i += 1

    doc.save(docx_path)
    size_kb = os.path.getsize(docx_path) / 1024
    print(f"  -> {os.path.basename(docx_path)} ({size_kb:.0f} KB)")


# ──────────────────────── Run ────────────────────────

if __name__ == '__main__':
    files = [
        ('STUDY_GUIDE_PERSON1_ML_AI.md',       'STUDY_GUIDE_PERSON1_ML_AI.docx'),
        ('STUDY_GUIDE_PERSON2_BACKEND.md',      'STUDY_GUIDE_PERSON2_BACKEND.docx'),
        ('STUDY_GUIDE_PERSON3_FRONTEND.md',     'STUDY_GUIDE_PERSON3_FRONTEND.docx'),
        ('STUDY_GUIDE_PERSON4_DB_DEPLOY.md',    'STUDY_GUIDE_PERSON4_DB_DEPLOY.docx'),
    ]

    print("Converting study guides to .docx ...\n")
    for md_name, docx_name in files:
        md_path = os.path.join(BASE_DIR, md_name)
        docx_path = os.path.join(BASE_DIR, docx_name)
        if not os.path.exists(md_path):
            print(f"  [SKIP] {md_name} not found")
            continue
        print(f"  Converting {md_name} ...")
        convert_md_to_docx(md_path, docx_path)

    print("\nDone! All .docx files saved in:")
    print(f"  {BASE_DIR}")
