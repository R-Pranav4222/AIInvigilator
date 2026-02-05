from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Title
title = doc.add_heading('AIInvigilator - Owner\'s Guide', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('Guide for Project Owner: Syncing and Pushing Changes')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(14)
subtitle_format.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()

intro = doc.add_paragraph('When teammates have made changes to your project and you want to continue working, follow these steps:')
intro_format = intro.runs[0]
intro_format.font.size = Pt(11)

doc.add_paragraph()

# Step 1
doc.add_heading('Step 1: Pull Latest Changes from GitHub', 1)
doc.add_paragraph('Before making any new changes, always sync your local repository with GitHub:')
p = doc.add_paragraph('cd E:\\witcher\\AIINVIGILATOR\\AIINVIGILATOR', style='Intense Quote')
p = doc.add_paragraph('', style='Intense Quote')
p = doc.add_paragraph('# Check which branch you\'re on', style='Intense Quote')
p = doc.add_paragraph('git branch', style='Intense Quote')
p = doc.add_paragraph('', style='Intense Quote')
p = doc.add_paragraph('# Switch to main branch (if not already there)', style='Intense Quote')
p = doc.add_paragraph('git checkout main', style='Intense Quote')
p = doc.add_paragraph('', style='Intense Quote')
p = doc.add_paragraph('# Pull all the latest changes from GitHub', style='Intense Quote')
p = doc.add_paragraph('git pull origin main', style='Intense Quote')
doc.add_paragraph('This downloads all the changes your teammates pushed and merges them into your local copy.')

# Step 2
doc.add_heading('Step 2: Check What Changed', 1)
p = doc.add_paragraph('# See the commit history', style='Intense Quote')
p = doc.add_paragraph('git log --oneline -10', style='Intense Quote')
p = doc.add_paragraph('', style='Intense Quote')
p = doc.add_paragraph('# See what files were changed', style='Intense Quote')
p = doc.add_paragraph('git log --stat -5', style='Intense Quote')
p = doc.add_paragraph('', style='Intense Quote')
p = doc.add_paragraph('# See who made changes', style='Intense Quote')
p = doc.add_paragraph('git log --pretty=format:"%h - %an, %ar : %s" -10', style='Intense Quote')

# Step 3
doc.add_heading('Step 3: Test the Updated Code', 1)
doc.add_paragraph('Always test after pulling changes to make sure everything works:')
p = doc.add_paragraph('python manage.py runserver', style='Intense Quote')
p = doc.add_paragraph('', style='Intense Quote')
p = doc.add_paragraph('# Test ML scripts if needed', style='Intense Quote')
p = doc.add_paragraph('cd ML', style='Intense Quote')
p = doc.add_paragraph('python front.py', style='Intense Quote')

# Step 4
doc.add_heading('Step 4: Make Your New Changes', 1)
doc.add_paragraph('Now you can safely make your changes:')
doc.add_paragraph('• Edit files', style='List Bullet')
doc.add_paragraph('• Add new features', style='List Bullet')
doc.add_paragraph('• Fix bugs', style='List Bullet')
doc.add_paragraph('• Update documentation', style='List Bullet')

# Step 5
doc.add_heading('Step 5: Check What You Changed', 1)
p = doc.add_paragraph('# See which files you modified', style='Intense Quote')
p = doc.add_paragraph('git status', style='Intense Quote')
p = doc.add_paragraph('', style='Intense Quote')
p = doc.add_paragraph('# See detailed changes', style='Intense Quote')
p = doc.add_paragraph('git diff', style='Intense Quote')
p = doc.add_paragraph('', style='Intense Quote')
p = doc.add_paragraph('# See changes in a specific file', style='Intense Quote')
p = doc.add_paragraph('git diff app/views.py', style='Intense Quote')

# Step 6
doc.add_heading('Step 6: Stage Your Changes', 1)
p = doc.add_paragraph('# Stage specific files', style='Intense Quote')
p = doc.add_paragraph('git add app/views.py', style='Intense Quote')
p = doc.add_paragraph('git add templates/malpractice_log.html', style='Intense Quote')
p = doc.add_paragraph('', style='Intense Quote')
p = doc.add_paragraph('# Or stage all modified files', style='Intense Quote')
p = doc.add_paragraph('git add .', style='Intense Quote')
p = doc.add_paragraph('', style='Intense Quote')
p = doc.add_paragraph('# To unstage a file if you added it by mistake', style='Intense Quote')
p = doc.add_paragraph('git reset HEAD filename.py', style='Intense Quote')

# Step 7
doc.add_heading('Step 7: Commit Your Changes', 1)
p = doc.add_paragraph('git commit -m "Improve detection accuracy for turning back feature"', style='Intense Quote')
doc.add_paragraph()
doc.add_paragraph('Examples of good commit messages:')
doc.add_paragraph('• "Fix video playback issue in malpractice log"', style='List Bullet')
doc.add_paragraph('• "Add export to PDF feature for reports"', style='List Bullet')
doc.add_paragraph('• "Update README with new installation steps"', style='List Bullet')
doc.add_paragraph('• "Refactor database queries for better performance"', style='List Bullet')

# Step 8
doc.add_heading('Step 8: Pull Again (Before Pushing)', 1)
important = doc.add_paragraph()
important.add_run('Important: ').bold = True
important.add_run('Pull one more time before pushing to avoid conflicts!')
doc.add_paragraph()
p = doc.add_paragraph('git pull origin main', style='Intense Quote')
doc.add_paragraph()
doc.add_paragraph('If there are NO conflicts:')
doc.add_paragraph('• Git will automatically merge', style='List Bullet')
doc.add_paragraph('• You\'re ready to push', style='List Bullet')
doc.add_paragraph()
doc.add_paragraph('If there ARE conflicts:')
doc.add_paragraph('1. Git will tell you which files have conflicts', style='List Number')
doc.add_paragraph('2. Open those files in VS Code', style='List Number')
doc.add_paragraph('3. Choose which code to keep or combine both', style='List Number')
doc.add_paragraph('4. Remove the conflict markers', style='List Number')
doc.add_paragraph('5. Save the file', style='List Number')
doc.add_paragraph('6. Stage and commit:', style='List Number')
p = doc.add_paragraph('git add .', style='Intense Quote')
p = doc.add_paragraph('git commit -m "Resolve merge conflicts"', style='Intense Quote')

# Step 9
doc.add_heading('Step 9: Push Your Changes to GitHub', 1)
p = doc.add_paragraph('git push origin main', style='Intense Quote')
doc.add_paragraph()
success = doc.add_paragraph('Your changes are now on GitHub! ✅')
success_format = success.runs[0]
success_format.font.color.rgb = RGBColor(0, 128, 0)
success_format.bold = True

# Add page break
doc.add_page_break()

# Alternative approach
doc.add_heading('Alternative: Work on Your Own Branch', 1)
doc.add_paragraph('(Recommended for Big Changes)')
doc.add_paragraph()
p = doc.add_paragraph('# Create and switch to a new branch', style='Intense Quote')
p = doc.add_paragraph('git checkout -b feature/my-new-feature', style='Intense Quote')
p = doc.add_paragraph('', style='Intense Quote')
p = doc.add_paragraph('# Make changes...', style='Intense Quote')
p = doc.add_paragraph('git add .', style='Intense Quote')
p = doc.add_paragraph('git commit -m "Add new feature"', style='Intense Quote')
p = doc.add_paragraph('', style='Intense Quote')
p = doc.add_paragraph('# Push your branch', style='Intense Quote')
p = doc.add_paragraph('git push origin feature/my-new-feature', style='Intense Quote')
p = doc.add_paragraph('', style='Intense Quote')
p = doc.add_paragraph('# Then create a Pull Request on GitHub and merge it yourself', style='Intense Quote')

# Quick Daily Workflow
doc.add_heading('Quick Daily Workflow for Owner', 1)

doc.add_heading('Morning: Start Work', 2)
p = doc.add_paragraph('cd E:\\witcher\\AIINVIGILATOR\\AIINVIGILATOR', style='Intense Quote')
p = doc.add_paragraph('git checkout main', style='Intense Quote')
p = doc.add_paragraph('git pull origin main', style='Intense Quote')
p = doc.add_paragraph('# ... make changes ...', style='Intense Quote')

doc.add_heading('Afternoon: Push Changes', 2)
p = doc.add_paragraph('git status', style='Intense Quote')
p = doc.add_paragraph('git add .', style='Intense Quote')
p = doc.add_paragraph('git commit -m "Description of changes"', style='Intense Quote')
p = doc.add_paragraph('git pull origin main  # Check for conflicts', style='Intense Quote')
p = doc.add_paragraph('git push origin main', style='Intense Quote')

doc.add_heading('Evening: If Teammates Pushed While You Were Working', 2)
p = doc.add_paragraph('git pull origin main  # Get their changes', style='Intense Quote')
p = doc.add_paragraph('# Test everything still works', style='Intense Quote')

# Add page break
doc.add_page_break()

# Common Scenarios
doc.add_heading('Common Scenarios', 1)

doc.add_heading('Scenario 1: Teammate Pushed While You Were Working', 2)
p = doc.add_paragraph('# You made changes but didn\'t commit yet', style='Intense Quote')
p = doc.add_paragraph('git stash  # Temporarily save your changes', style='Intense Quote')
p = doc.add_paragraph('git pull origin main  # Get teammate\'s changes', style='Intense Quote')
p = doc.add_paragraph('git stash pop  # Restore your changes on top', style='Intense Quote')
p = doc.add_paragraph('# Resolve any conflicts if needed', style='Intense Quote')
p = doc.add_paragraph('git add .', style='Intense Quote')
p = doc.add_paragraph('git commit -m "Your changes"', style='Intense Quote')
p = doc.add_paragraph('git push origin main', style='Intense Quote')

doc.add_heading('Scenario 2: You Committed but Teammate Also Pushed', 2)
p = doc.add_paragraph('# You committed locally', style='Intense Quote')
p = doc.add_paragraph('git pull origin main  # Get teammate\'s changes', style='Intense Quote')
p = doc.add_paragraph('# Git will try to auto-merge', style='Intense Quote')
p = doc.add_paragraph('# If conflicts, resolve them', style='Intense Quote')
p = doc.add_paragraph('git add .', style='Intense Quote')
p = doc.add_paragraph('git commit -m "Resolve conflicts"', style='Intense Quote')
p = doc.add_paragraph('git push origin main', style='Intense Quote')

doc.add_heading('Scenario 3: You Want to See What Changed Before Pulling', 2)
p = doc.add_paragraph('# Fetch changes without merging', style='Intense Quote')
p = doc.add_paragraph('git fetch origin', style='Intense Quote')
p = doc.add_paragraph('', style='Intense Quote')
p = doc.add_paragraph('# Compare your local with remote', style='Intense Quote')
p = doc.add_paragraph('git diff main origin/main', style='Intense Quote')
p = doc.add_paragraph('', style='Intense Quote')
p = doc.add_paragraph('# See commits you don\'t have yet', style='Intense Quote')
p = doc.add_paragraph('git log main..origin/main', style='Intense Quote')
p = doc.add_paragraph('', style='Intense Quote')
p = doc.add_paragraph('# If everything looks good, merge', style='Intense Quote')
p = doc.add_paragraph('git merge origin/main', style='Intense Quote')

# Add page break
doc.add_page_break()

# Important Tips
doc.add_heading('Important Tips for Owner', 1)
doc.add_paragraph('1. Always git pull before starting work - Get latest changes first', style='List Number')
doc.add_paragraph('2. Pull again before pushing - Avoid conflicts', style='List Number')
doc.add_paragraph('3. Commit frequently - Don\'t accumulate too many changes', style='List Number')
doc.add_paragraph('4. Test after pulling - Make sure teammates\' changes work', style='List Number')
doc.add_paragraph('5. Write clear commit messages - Help your team understand what changed', style='List Number')
doc.add_paragraph('6. Communicate - Let team know if you\'re working on same files', style='List Number')
doc.add_paragraph('7. Use branches for big features - Keep main branch stable', style='List Number')
doc.add_paragraph('8. Review Pull Requests promptly - Don\'t block your teammates', style='List Number')

# Check Repository Status
doc.add_heading('Check Repository Status Anytime', 1)
p = doc.add_paragraph('# See if you\'re behind/ahead of GitHub', style='Intense Quote')
p = doc.add_paragraph('git status', style='Intense Quote')
p = doc.add_paragraph('', style='Intense Quote')
p = doc.add_paragraph('# See remote repository info', style='Intense Quote')
p = doc.add_paragraph('git remote -v', style='Intense Quote')
p = doc.add_paragraph('', style='Intense Quote')
p = doc.add_paragraph('# See all branches (local and remote)', style='Intense Quote')
p = doc.add_paragraph('git branch -a', style='Intense Quote')
p = doc.add_paragraph('', style='Intense Quote')
p = doc.add_paragraph('# See who\'s working on what', style='Intense Quote')
p = doc.add_paragraph('git log --all --graph --oneline --decorate', style='Intense Quote')

# Emergency Commands
doc.add_heading('Emergency: Undo Changes', 1)
doc.add_paragraph('Undo changes to a specific file (not committed yet):')
p = doc.add_paragraph('git checkout -- filename.py', style='Intense Quote')
doc.add_paragraph()

doc.add_paragraph('Undo all uncommitted changes:')
p = doc.add_paragraph('git reset --hard', style='Intense Quote')
doc.add_paragraph()

doc.add_paragraph('Undo last commit (keep changes):')
p = doc.add_paragraph('git reset --soft HEAD~1', style='Intense Quote')
doc.add_paragraph()

doc.add_paragraph('Undo last commit (discard changes):')
p = doc.add_paragraph('git reset --hard HEAD~1', style='Intense Quote')
doc.add_paragraph()

doc.add_paragraph('Restore to a specific commit:')
p = doc.add_paragraph('git reset --hard commit-hash', style='Intense Quote')

# Visual Summary
doc.add_heading('Visual Summary: Your Workflow as Owner', 1)
workflow = doc.add_paragraph()
workflow.add_run('1. git pull origin main').bold = True
workflow.add_run(' ← Get latest changes\n')
workflow.add_run('2. Make your changes').bold = True
workflow.add_run(' ← Edit files\n')
workflow.add_run('3. git add .').bold = True
workflow.add_run(' ← Stage changes\n')
workflow.add_run('4. git commit -m "message"').bold = True
workflow.add_run(' ← Commit changes\n')
workflow.add_run('5. git pull origin main (again!)').bold = True
workflow.add_run(' ← Check for conflicts\n')
workflow.add_run('6. git push origin main').bold = True
workflow.add_run(' ← Upload to GitHub')

# Footer
doc.add_paragraph()
doc.add_paragraph()
footer = doc.add_paragraph('Key Takeaway: As the owner, your workflow is the same as contributors, but you push directly to main. Always pull before you start and pull before you push! 🚀')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_format = footer.runs[0]
footer_format.font.size = Pt(12)
footer_format.font.color.rgb = RGBColor(0, 100, 200)
footer_format.bold = True

# Save document
doc.save('GitHub_Owner_Guide.docx')
print("Word document created successfully: GitHub_Owner_Guide.docx")
